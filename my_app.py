from docx import Document
from docx.shared import Inches

document = Document()

# foto de perfil
document.add_picture('Foto de Perfil.jpg', width=Inches(1.4))

# dados pessoais
nome = input('Qual é seu nome? ')
celular = input('Qual seu número de celular? ')
email = input ('Qual seu e-mail? ')
idade = input ('Quantos anos você tem? ')

document.add_paragraph(
    nome + '    -    ' + idade + ' anos' + '    -    ' + '+55 11 ' + celular + '    -    ' + email)

# perfil
document.add_heading('Sobre mim')
sobre_mim = input('Me conte sobre você! ')
document.add_paragraph(sobre_mim)

# experiências
document.add_heading('Experiências')
p = document.add_paragraph()

empresa = input('Qual empresa? ')
data_inicial = input('Quando começou lá? ')
data_final = input('Quando saiu de lá? ')

p.add_run(empresa + ' ').bold = True
p.add_run('       ' + data_inicial + ' - ' + data_final + '\n').italic = True

detalhes = input('Quais foram suas experiências na ' + empresa + '? ')
p.add_run(detalhes)

# demais experiências
while True:
    mais_experiencias = input(
        'Trabalhou em mais algum lugar? ')
    if mais_experiencias.lower() == 'sim':
        p = document.add_paragraph()

        empresa = input('Qual empresa? ')
        data_inicial = input('Quando começou lá? ')
        data_final = input('Quando saiu de lá? ')

        p.add_run(empresa + ' ').bold = True
        p.add_run('       ' + data_inicial + ' - ' + data_final + '\n').italic = True

        detalhes = input('Quais foram suas experiências na ' + empresa + ' ? ')
        p.add_run(detalhes)
    else:
        break


# habilidades
document.add_heading('Habilidades')
h = document.add_paragraph()

habilidade = input('Qual sua habilidade? ')
nivel = input('Qual seu nível de experiência em ' + habilidade + '? (Avaliar de 1~5) ')

h.add_run(habilidade + ' ').bold = True
h.add_run(' - ' + nivel + ' (1~5)' + '\n').italic = True

# demais habilidades
while True:
    mais_habilidades = input(
        'Possui mais habilidades? ')
    if mais_habilidades.lower() == 'sim':
        h = document.add_paragraph()
    
        habilidade = input('Qual sua habilidade? ')
        nivel = input('Qual seu nível de experiência em ' + habilidade + '? (Avaliar de 1~5) ')

        h.add_run(habilidade + ' ').bold = True
        h.add_run(' - ' + nivel + ' (1~5)' + '\n').italic = True
    else:
        break


document.save('cv.docx') 